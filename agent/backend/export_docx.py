"""把 Markdown / LaTeX 文本导出为 Word(.docx)。

优先用 pandoc（pypandoc）转换——能较好地处理公式、标题、列表、表格；
本机没装 pandoc 时回退到 python-docx 做基础排版（标题/段落/列表/代码块，公式以文本保留）。
"""
import io
import re


def to_docx(content: str, source_fmt: str = "markdown") -> bytes:
    """content：文档源文本；source_fmt：'markdown' 或 'latex'。返回 .docx 字节。"""
    fmt = "latex" if source_fmt in ("latex", "tex") else "markdown"
    data = _try_pandoc(content, fmt)
    if data is not None:
        return data
    # 回退：LaTeX 先粗略剥成纯文本，再按 Markdown 走基础排版
    text = _latex_to_text(content) if fmt == "latex" else content
    return _markdown_to_docx(text)


def _try_pandoc(content: str, fmt: str) -> bytes | None:
    try:
        import pypandoc
    except Exception:
        return None
    import tempfile
    import os
    tmp = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = f.name
        pypandoc.convert_text(content, "docx", format=fmt, outputfile=tmp)
        with open(tmp, "rb") as f:
            return f.read()
    except Exception:
        return None
    finally:
        if tmp and os.path.exists(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass


def _markdown_to_docx(md: str) -> bytes:
    """python-docx 基础排版：标题、段落、无序/有序列表、围栏代码块。公式按文本保留。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    lines = md.splitlines()
    in_code = False
    code_buf: list[str] = []

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        code_buf.clear()

    for raw in lines:
        line = raw.rstrip("\n")
        fence = line.strip().startswith("```")
        if fence:
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_strip_inline(m.group(2)), level=level)
            continue
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if m:
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Bullet")
            continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            doc.add_paragraph(_strip_inline(m.group(1)), style="List Number")
            continue
        if not line.strip():
            continue
        doc.add_paragraph(_strip_inline(line))

    flush_code()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_inline(text: str) -> str:
    """去掉常见 Markdown 行内标记（**、*、`、[]() ），公式 $...$ 原样保留为文本。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
    return text


def _latex_to_text(tex: str) -> str:
    """pandoc 不可用时，把 LaTeX 粗略剥成可读文本（仅用于回退）。"""
    # 取 document 环境内容（若有）
    m = re.search(r"\\begin\{document\}(.*?)\\end\{document\}", tex, re.S)
    body = m.group(1) if m else tex
    body = re.sub(r"\\(section|subsection|subsubsection)\*?\{(.*?)\}", r"\n# \2\n", body)
    body = re.sub(r"\\(title|author|date)\{(.*?)\}", r"\2", body)
    body = re.sub(r"\\item\s*", "- ", body)
    body = re.sub(r"\\begin\{[^}]+\}|\\end\{[^}]+\}", "", body)
    body = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{.*?\})?", "", body)
    body = re.sub(r"[{}]", "", body)
    body = re.sub(r"%.*", "", body)
    body = re.sub(r"\n{3,}", "\n\n", body)
    return body.strip()
